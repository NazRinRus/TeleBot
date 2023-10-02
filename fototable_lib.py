import json
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from config import keys
from telebot import types
from datetime import datetime

# Класс принимающий данные объекта фототаблицы, посредствам json-файла, формирующий docx-файл
class FototableDocx:
    def __init__(self, id_user, kusp):
        self.id_user = id_user
        self.kusp = kusp

    # извлечение данных фототаблицы с файла json
    def fototable_json_open(self):
        with open(f'files/{self.id_user}/{self.kusp}/{self.kusp}.json') as f:
            file_content = f.read()
            templates = json.loads(file_content)
        return templates

    # извлечение данных пользователя с файла json
    def user_json_open(self):
        with open('users.json') as f:
            file_content = f.read()
            templates = json.loads(file_content)
        return templates

    # блок работы с шаблоном фототаблицы
    def fototable_docx_save(self):
        templates = self.fototable_json_open()
        id_user = templates['id_user']
        kusp = templates['kusp']
        date_kusp = templates['current_datetime']
        kusp_info = f'{kusp} от {date_kusp}'
        images = templates['images']
        user = self.user_json_open() # это список пользователей, не конкретный пользователь

        doc = docx.Document('files/fototable_sample.docx')  # создаю экземпляр шаблона фототаблицы

        # верхний колонтитул
        header = doc.sections[0].header # получаю объект верхнего колонтитула
        title = header.tables[0]  # таблица с информацией о КУСП (номер КУСП и дата)
        kusp_cell = title.cell(0, 1)  # ячейка содержащая информацию о КУСП
        kusp_cell.text = kusp_info  # текст ячейки содержащей номер КУСП

        # Содержание фототаблицы
        number_of_rows = len(images) * 2  # количество строк в таблице равно количеству фотографий умноженное на 2 (строка с пояснениями)
        for row in range(number_of_rows):
            if (row % 2) == 0:
                # Добавляем пустой абзац
                p = doc.add_paragraph()
                # Добавляем пустой прогон
                run = p.add_run()
                run.add_picture(images[row // 2]["name"], width=docx.shared.Cm(12))
                # выравниваем картинку посередине страницы
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # cell.add_picture('files/380009616/3124/file_12.jpg', width=docx.shared.Cm(12))
                # images[row // 2]["name"]
            else:
                # Добавляем абзац с порядковым номером фотографии
                p = doc.add_paragraph(f'Фото №{str(row // 2 + 1)}. ')
                # Добавляем прогон с коментарием к фотографии
                run = p.add_run(images[row // 2]["photo_description"])
                # выравниваем коментарий по ширине страницы
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Добавляем пустой абзац для визуального разделения
                doc.add_paragraph()

        # нижний колонтитул
        footer = doc.sections[0].footer  # получаю объект нижнего колонтитула
        title_footer = footer.tables[0]  # таблица с информацией о должностном лице
        job_title_cell = title_footer.cell(1, 1)  # ячейка содержащая информацию о должности
        job_title_cell.text = user[str(id_user)]['job_title']  # текст ячейки содержащей должность
        rank_cell = title_footer.cell(3, 0)  # ячейка содержащая информацию о звании
        rank_cell.text = user[str(id_user)]['rank']  # текст ячейки содержащей звание
        fio_cell = title_footer.cell(3, 2)  # ячейка содержащая фамилию и инициалы
        fio_cell.text = user[str(id_user)]['fio']  # текст ячейки содержащей фамилию и инициалы
        fio_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # сохранение файла в папке с фотографиями под именем КУСП
        doc.save(f'files/{id_user}/{kusp}/{kusp}.docx')

if __name__ == '__main__':

    fototabledocx = FototableDocx(380009616, 3124)
    fototabledocx.fototable_docx_save()
